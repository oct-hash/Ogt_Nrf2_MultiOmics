"""Generate STAR Methods section as docx"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from pathlib import Path

doc = Document()
style = doc.styles["Normal"]
style.font.size = Pt(11)
style.font.name = "Arial"
style.paragraph_format.space_after = Pt(4)

def heading(text, level=1):
    doc.add_heading(text, level=level)

def subheading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)

def body(text):
    doc.add_paragraph(text)

# ================================================================
heading("STAR Methods", level=1)

# ── Resource Availability ───────────────────────────────────────
heading("Resource Availability", level=2)

subheading("Lead Contact")
body("Further information and requests for resources and reagents should be directed to "
     "and will be fulfilled by the Lead Contact, [NAME] ([EMAIL]).")

subheading("Materials Availability")
body("This study did not generate new unique reagents.")

subheading("Data and Code Availability")
body("All raw and processed sequencing data have been deposited in the Gene Expression Omnibus (GEO) "
     "under accession numbers: GSE193997, GSE307385 (bulk RNA-seq); GSE240848 (snRNA-seq); "
     "GSE163129 (scRNA-seq). Proteomics data have been deposited to the ProteomeXchange Consortium "
     "via the PRIDE partner repository: PXD035154, PXD040135, PXD046631. "
     "The external validation dataset GSE57338 is publicly available at GEO. "
     "GTEx v8 data were accessed via the GTEx Portal (https://gtexportal.org). "
     "All original code has been deposited at GitHub ([URL]) and is publicly available as of the date of publication. "
     "Any additional information required to reanalyze the data reported in this paper is available from the Lead Contact upon request.")

# ── Experimental Model ──────────────────────────────────────────
heading("Experimental Model and Subject Details", level=2)

subheading("Mouse Model of Myocardial Ischemia-Reperfusion")
body("[USER: Fill in — mouse strain, age, sex, I/R protocol details (LAD ligation duration, reperfusion time), "
     "sham procedure, group sizes, time points. This section should describe the animal model used for "
     "GSE193997, GSE307385, and the proteomics datasets.]")

subheading("Human Heart Failure Samples (GSE57338)")
body("The external validation dataset GSE57338 consists of Affymetrix Human Gene 1.0 ST Array "
     "(GPL11532) expression profiles from left ventricular myocardial samples of 313 individuals, "
     "including 177 with heart failure (idiopathic dilated cardiomyopathy or ischemic cardiomyopathy) "
     "and 136 non-failing controls. All samples were obtained with informed consent and institutional "
     "review board approval as described in the original publication (Liu et al., 2015, PMID: 25528681).")

subheading("GTEx v8 Healthy Controls")
body("Median gene-level TPM values for Heart - Left Ventricle (n = 431) were obtained from the "
     "Genotype-Tissue Expression (GTEx) project v8 release via the GTEx Portal API v2 "
     "(https://gtexportal.org). GTEx donor characteristics and tissue procurement procedures "
     "are described in the GTEx consortium publications.")

# ── Method Details ──────────────────────────────────────────────
heading("Method Details", level=2)

subheading("Bulk RNA-seq Processing and Differential Expression Analysis")
body("Raw RNA-seq reads from GSE193997 and GSE307385 were aligned to the mouse reference genome "
     "(GRCm38/mm10) using STAR v2.7. Gene-level counts were quantified with featureCounts. "
     "Differential expression analysis was performed using DESeq2 (v1.34). "
     "Genes with adjusted P-value (Benjamini-Hochberg) < 0.05 were considered differentially expressed. "
     "For the primary discovery analysis in GSE193997 (I/R 6h vs Sham), 6,556 DEGs met this threshold. "
     "[USER: add details about GSE307385 integration — meta-analysis method, batch correction, etc.]")

subheading("Gene Set Enrichment Analysis (GSEA)")
body("GSEA was performed using the fgsea R package (v1.20) with genes ranked by signed "
     "-log10(P-value) x log2FoldChange. The Nrf2-ARE gene set was curated from literature "
     "[USER: cite source of gene set], and the Ferroptosis gene set was obtained from FerrDb. "
     "The Hallmark gene set collection from MSigDB was used for broad pathway annotation. "
     "Permutation testing (10,000 permutations) was used to assess significance.")

subheading("Single-Nucleus RNA-seq Analysis")
body("The snRNA-seq dataset GSE240848 was processed using Cell Ranger (10x Genomics) and "
     "downstream analysis was performed in Seurat v5 / Scanpy v1.9. "
     "Quality control removed nuclei with >10% mitochondrial reads, <200 or >5,000 detected genes. "
     "After QC, 48,633 nuclei were retained. Normalization was performed using SCTransform. "
     "Dimensionality reduction was performed by PCA followed by UMAP. "
     "Cell types were annotated using canonical marker genes: "
     "Cardiomyocyte (MYH6, TNNT2), Fibroblast (COL1A1, DCN), Endothelial (PECAM1, CDH5), "
     "Macrophage (CD68, CD163), Neutrophil (S100A8, S100A9), "
     "T Cell (CD3D, CD3E), B Cell (CD79A, MS4A1), "
     "Pericyte (PDGFRB, RGS5), Smooth Muscle Cell (ACTA2, MYH11). "
     "[USER: verify marker genes match your annotation.]")

subheading("Cell-Cell Communication Inference")
body("Cell-cell communication was inferred using two independent methods: CellChat (v1.6) "
     "and CellPhoneDB (v4, via the LIANA framework). Both methods were applied to the "
     "normalized snRNA-seq count matrix with the same cell type annotations. "
     "For CellChat, the default ligand-receptor database (CellChatDB.mouse) was used, "
     "and significant interactions were identified by permutation test (P < 0.05). "
     "For CellPhoneDB, significance was assessed using the built-in permutation framework. "
     "Consensus interactions were defined as ligand-receptor pairs identified as significant "
     "by both methods with matching sender and receiver cell types. "
     "Analysis focused on signaling from myeloid cells (Macrophage, Neutrophil) to "
     "Cardiomyocyte and Pericyte populations.")

subheading("Proteomics Data Processing")
body("Three independent proteomics datasets were analyzed: PXD035154 (label-free quantification, "
     "mouse cardiac I/R, n = 3 per group, 2 fractionated pools), "
     "PXD040135 (label-free quantification, mouse cardiac I/R, n = 3 per group), "
     "and PXD046631 (label-free quantification, mouse cardiac I/R 6h, n = 4 per group). "
     "[USER: confirm sample sizes and experimental details.] "
     "MaxQuant output (proteinGroups.txt) was processed to remove contaminants and reverse hits. "
     "Protein intensities were log2-transformed. For cross-platform comparison, "
     "gene symbols were mapped to human orthologs via the mygene Python package using HomoloGene. "
     "Proteins detected in at least 2 of 3 replicates per group were retained for analysis.")

subheading("mRNA-Protein Discordance Analysis (Residual Method)")
body("For the PXD046631 dataset (2,308 genes with matched mRNA and protein measurements), "
     "mRNA-protein discordance was quantified by linear regression of log2(protein fold change) "
     "against log2(mRNA fold change). The regression yielded R = 0.253 (slope = 0.144), "
     "indicating weak genome-wide mRNA-protein correlation. For each gene, the Studentized "
     "residual from this regression was computed. Positive residuals indicate protein abundance "
     "exceeding mRNA-based predictions (suggesting translational activation or protein stabilization); "
     "negative residuals indicate lower-than-expected protein (suggesting translational repression "
     "or protein degradation). Directional concordance was defined as agreement in the sign of "
     "mRNA and protein log2FC (both positive or both negative).")

subheading("OGT Substrate Enrichment Analysis")
body("High-stringency human OGT protein interaction partners (n = 782) were obtained from "
     "OGT-PIN v2.0 (Ma et al., 2021; https://oglcnac.org/ogt-pin/), a curated database of "
     "experimentally identified OGT-interacting proteins. Mouse DEGs were mapped to human "
     "orthologs by case-insensitive matching (mouse gene symbols uppercased to human equivalents). "
     "Enrichment of OGT interactors among DEGs was assessed by Fisher's exact test "
     "(background: 20,000 human protein-coding genes). Functional categorization of network "
     "genes (O-GlcNAc Cycling, NRF2 Pathway, Ferroptosis Core, Antioxidant Targets, Signaling) "
     "was performed to assess pathway-specific overlap patterns.")

subheading("External Validation Using GTEx and GSE57338")
body("GTEx v8 median TPM values for Heart - Left Ventricle (n = 431 healthy donors) were "
     "queried via the GTEx Portal API v2 for core genes: OGT (ENSG00000147162.13), "
     "NFE2L2 (ENSG00000116044.15), HMOX1 (ENSG00000100292.16), "
     "SLC7A11 (ENSG00000151012.13), and BACH1 (ENSG00000156273.15). "
     "For GSE57338 (313 human left ventricular samples), Affymetrix probe-to-gene mapping "
     "was performed using the GPL11532 platform annotation. Differential expression between "
     "177 HF and 136 NF samples was assessed by Welch's t-test. Directional consistency with "
     "the GSE193997 discovery DEGs was evaluated by binomial test (null hypothesis: 50% "
     "same direction). For cross-platform comparison, mouse gene symbols from GSE193997 "
     "were matched to human symbols (case-insensitive) on the GSE57338 array (5,514 common genes).")

subheading("WGCNA and Network Analysis")
body("[USER: Fill in — describe WGCNA parameters (soft threshold, network type, "
     "module detection method, module-trait correlation). Include Boolean network "
     "modeling description if retained in Paper 1, or note that it is described in "
     "a companion paper.]")

# ── Quantification and Statistical Analysis ─────────────────────
heading("Quantification and Statistical Analysis", level=2)

body("All statistical analyses were performed in R (v4.2) and Python (v3.10). "
     "Key software packages and versions are listed in the Key Resources Table. "
     "Statistical details for each analysis are provided in the corresponding Method Details "
     "subsection and Figure Legends. Briefly:")

stats = [
    "Bulk RNA-seq differential expression: DESeq2 with Benjamini-Hochberg FDR correction; significance threshold: adjusted P < 0.05.",
    "GSEA: fgsea with 10,000 permutations; significance threshold: FDR < 0.25 (standard GSEA threshold).",
    "Cell-cell communication: CellChat permutation test (P < 0.05); CellPhoneDB built-in permutation (P < 0.05); consensus defined by overlap of both methods.",
    "mRNA-protein residual analysis: linear regression (scipy.stats.linregress); Studentized residuals; concordance defined as same-sign log2FC.",
    "OGT substrate enrichment: Fisher's exact test (one-sided), background = 20,000 human protein-coding genes.",
    "Directional consistency (GSE57338): binomial test (one-sided) against null probability of 0.5.",
    "Cross-platform comparisons used direction of change (log2FC sign) rather than absolute effect sizes, given differences in measurement scales between RNA-seq and microarray platforms.",
    "[USER: add mouse group sizes and statistical tests used for animal experiments.]",
]
for s in stats:
    doc.add_paragraph(s, style="List Bullet")

body("Sample sizes (n) and statistical tests are reported in each Figure Legend. "
     "No statistical methods were used to predetermine sample size. "
     "Experiments were not randomized, and investigators were not blinded to "
     "allocation during experiments and outcome assessment. "
     "[USER: modify if randomization/blinding were used.]")

# ── Key Resources Table ─────────────────────────────────────────
heading("Key Resources Table", level=2)

table_items = [
    ("Deposited Data", "", ""),
    ("Bulk RNA-seq (mouse I/R)", "GEO: GSE193997", "This paper; [USER: add citation if published]"),
    ("Bulk RNA-seq (mouse I/R time-series)", "GEO: GSE307385", "This paper"),
    ("snRNA-seq (mouse I/R heart)", "GEO: GSE240848", "This paper"),
    ("scRNA-seq (mouse I/R heart)", "GEO: GSE163129", "This paper"),
    ("Proteomics (mouse I/R, LFQ pool 1)", "PRIDE: PXD035154", "This paper"),
    ("Proteomics (mouse I/R, LFQ pool 2)", "PRIDE: PXD040135", "This paper"),
    ("Proteomics (mouse I/R, LFQ)", "PRIDE: PXD046631", "This paper"),
    ("External validation (human HF)", "GEO: GSE57338", "Liu et al., 2015, PMID: 25528681"),
    ("Healthy donor expression", "GTEx v8", "GTEx Consortium, 2020"),
    ("OGT interactome", "OGT-PIN v2.0", "Ma et al., 2021, PMID: 34502531"),
    ("GPL11532 annotation", "GEO: GPL11532", "Affymetrix Human Gene 1.0 ST"),
    ("", "", ""),
    ("Software and Algorithms", "", ""),
    ("STAR", "v2.7", "https://github.com/alexdobin/STAR"),
    ("featureCounts", "v2.0", "http://subread.sourceforge.net/"),
    ("DESeq2", "v1.34", "https://bioconductor.org/packages/DESeq2/"),
    ("fgsea", "v1.20", "https://bioconductor.org/packages/fgsea/"),
    ("Cell Ranger", "v7", "10x Genomics"),
    ("Seurat", "v5", "https://satijalab.org/seurat/"),
    ("Scanpy", "v1.9", "https://scanpy.readthedocs.io/"),
    ("CellChat", "v1.6", "https://github.com/sqjin/CellChat"),
    ("CellPhoneDB", "v4", "https://github.com/ventolab/CellphoneDB"),
    ("LIANA", "v1.0", "https://github.com/saezlab/liana"),
    ("MaxQuant", "v2.0", "https://www.maxquant.org/"),
    ("Python", "v3.10", "https://www.python.org/"),
    ("R", "v4.2", "https://www.r-project.org/"),
    ("Custom analysis code", "GitHub: [URL]", "This paper"),
]

table = doc.add_table(rows=1, cols=3)
table.style = "Light List Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "REAGENT or RESOURCE"; hdr[1].text = "SOURCE"; hdr[2].text = "IDENTIFIER"
for p in hdr:
    for run in p.paragraphs[0].runs: run.bold = True

for item in table_items:
    row = table.add_row()
    for i, text in enumerate(item):
        row.cells[i].text = text

# ── Save ──────────────────────────────────────────────────────────
out = Path(r"D:\Cardiac_Ogt_MultiOmics\manuscript\STAR_Methods_draft.docx")
doc.save(str(out))
print(f"Saved: {out}")
