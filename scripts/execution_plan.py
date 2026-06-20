"""Generate the master execution plan for Paper 1 writing"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.size = Pt(11)
style.font.name = "Arial"
style.paragraph_format.space_after = Pt(4)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("论文一：分步执行计划（含验证门槛）")
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph("目标: 完成一篇 Cell Reports 级多组学发现论文，7张主图，≤7000字")
doc.add_paragraph("核心论点: 急性心肌I/R中，Ogt下调通过RELA/NF-kB间接激活Nrf2，髓系SPP1-CD44旁分泌放大铁死亡信号，铁蛋白自噬加剧心肌损伤。该机制是急性损伤阶段特异性的。")
doc.add_paragraph("")

# ============================================================
doc.add_heading("阶段一：数据与代码就绪 (Day 1-2)", level=1)

steps = [
    {
        "step": "1.1 分析脚本归档",
        "input": "scripts/*.py (step1-4 分析脚本)",
        "action": [
            "清理分析脚本中的临时路径和调试代码",
            "统一文件路径为相对路径 (基于项目根目录)",
            "添加 argparse 支持，可从命令行调用",
            "每个脚本添加 docstring 说明输入/输出/依赖",
        ],
        "output": "scripts/ 下每个 .py 文件可独立运行，python stepX_xxx.py 即可复现",
        "verify": "在新终端中逐个运行每个 step 脚本，确认无报错且输出文件与之前一致",
    },
    {
        "step": "1.2 创建 GitHub 仓库",
        "input": "清理后的脚本 + 输出文件",
        "action": [
            "创建 README.md (研究概述 + 环境配置 + 运行说明)",
            "创建 environment.yml 或 requirements.txt",
            "上传 scripts/ + output_tables/ (不含大数据文件)",
            "STAR Methods 中引用 GitHub URL",
        ],
        "output": "github.com/<user>/Ogt_Nrf2_MultiOmics (公开或审稿期间私密)",
        "verify": "另一台机器上 git clone 后能按 README 复现核心分析",
    },
    {
        "step": "1.3 数据可访问性确认",
        "input": "所有GEO/PRIDE accession numbers",
        "action": [
            "确认 GSE193997, GSE240848, GSE163129, GSE307385 在 GEO 中为公开状态",
            "确认 PXD035154, PXD040135, PXD053525 在 ProteomeXchange 中为公开状态",
            "确认 GSE57338 (外部验证) 为公开状态",
            "在 Data Availability 段落中列出所有 accession",
        ],
        "output": "完整的 Data Availability 段落，每个数据集都有可访问的 accession number",
        "verify": "用浏览器逐一打开每个 accession URL，确认可公开访问",
    },
]

for s in steps:
    doc.add_heading(s["step"], level=2)
    p = doc.add_paragraph(); run = p.add_run("输入: "); run.bold = True; p.add_run(s["input"])
    p = doc.add_paragraph(); run = p.add_run("操作: "); run.bold = True
    for a in s["action"]: doc.add_paragraph(a, style="List Bullet")
    p = doc.add_paragraph(); run = p.add_run("输出: "); run.bold = True; p.add_run(s["output"])
    p = doc.add_paragraph(); run = p.add_run("验证: "); run.bold = True; p.add_run(s["verify"])
    doc.add_paragraph("")

# ============================================================
doc.add_heading("阶段二：图表制作 (Day 3-7)", level=1)

fig_steps = [
    {
        "step": "2.1 Fig 1: 研究概览",
        "target": "审稿人2分钟内理解研究规模与设计",
        "action": [
            "Panel A: 设计示意图 (BioRender或draw.io)，展示 2 bulk + 1 sn + 1 sc + 3 prot 的数据整合架构",
            "Panel B: 火山图 (6,777 DEGs)，标注 Ogt, Hmox1, Slc7a11, Rela",
        ],
        "verify": "给非本领域同事看图1，2分钟后问'这个研究做了什么'，能正确回答即通过",
    },
    {
        "step": "2.2 Fig 2: Nrf2/铁死亡通路 + GTEx基线",
        "target": "通路协同富集 + 健康开关模式清晰可见",
        "action": [
            "Panel A: GSEA enrichment plot (Nrf2-ARE, NES=1.77, FDR=0.001)",
            "Panel B: 热图 (top DEGs in Nrf2 + ferroptosis pathways)",
            "Panel C: GTEx v8 箱线图/柱状图 (Ogt 39 vs Hmox1 3 TPM, n=431)",
        ],
        "verify": "三个Panel各自传达一个清晰信息，Panel C的开关对比一目了然",
    },
    {
        "step": "2.3 Fig 3: mRNA-蛋白偏离",
        "target": "从50%宏观一致性落地为FTH1/TFRC具体机制",
        "action": [
            "Panel A: 散点图 log2(Prot) vs log2(RNA)，R=0.253，标注FTH1/TFRC/OGT",
            "Panel B: 残差瀑布图或排序图",
            "Panel C: FTH1和TFRC的RNA vs 蛋白方向对比放大图",
        ],
        "verify": "读者能在10秒内找到FTH1和TFRC并理解它们是outlier",
    },
    {
        "step": "2.4 Fig 4: snRNA-seq 细胞二分",
        "target": "一眼看出 Ogt在心肌 vs Nrf2靶基因在髓系",
        "action": [
            "Panel A: UMAP (9种细胞类型，48,633核)",
            "Panel B: Feature Plot — Ogt (心肌+周细胞高亮)",
            "Panel C: Feature Plot — Hmox1/Slc7a11 (髓系高亮)",
            "Panel D: 关键基因小提琴图 (Ogt, Hmox1, Slc7a11, Rela by cell type)",
        ],
        "verify": "Feature plots的颜色对比直观强烈，不需要读图注就能理解空间分离",
    },
    {
        "step": "2.5 Fig 5: SPP1-CD44 旁分泌",
        "target": "双方法确认 + SPP1-CD44是核心信号",
        "action": [
            "Panel A: CellChat circle plot (髓系->心肌/周细胞连线高亮)",
            "Panel B: Upset图 (CellChat x CellPhoneDB交集=8对)",
            "Panel C: SPP1-CD44信号模式图 (配体来源->受体->下游Nrf2靶基因)",
        ],
        "verify": "Panel C模式图能独立讲清楚 SPP1-CD44 如何连接旁分泌与铁死亡",
    },
    {
        "step": "2.6 Fig 6: OGT底物网络 + RELA中介",
        "target": "OGT不直接接触Nrf2 -> RELA是桥梁",
        "action": [
            "Panel A: OGT-PIN x DEG 富集条形图 (OR=2.02, P<0.001)",
            "Panel B: 27节点布尔网络中标注 OGT互作关系 (OGT, OGA, RELA 高亮)",
            "Panel C: Ogt -> RELA -> Nrf2 -> Ferroptosis 模型图",
        ],
        "verify": "Panel C的模型图是全文最重要的机制总结，应能独立作为Graphical Abstract",
    },
    {
        "step": "2.7 Fig 7: 外部验证 + 整合总结",
        "target": "诚实报告阶段特异性 -> 科学叙事升级",
        "action": [
            "Panel A: GSE57338方向一致性散点图 (x=your logFC, y=GSE57338 logFC, 标HMOX1/SLC7A11反向)",
            "Panel B: 急性I/R vs 慢性心衰的模式对比图",
            "Panel C: Graphical Abstract (Ogt -> RELA -> Nrf2 -> Ferroptosis, 急性I/R特异)",
        ],
        "verify": "读者看完Panel B后理解'这个发现不是泛化的心衰标志物，而是急性损伤反应'",
    },
]

for s in fig_steps:
    doc.add_heading(s["step"], level=2)
    p = doc.add_paragraph(); run = p.add_run("目标: "); run.bold = True; p.add_run(s["target"])
    for a in s["action"]: doc.add_paragraph(a, style="List Bullet")
    p = doc.add_paragraph(); run = p.add_run("验证: "); run.bold = True; p.add_run(s["verify"])
    doc.add_paragraph("")

# Supplementary figures
doc.add_heading("2.8 Supplementary Figures (建议10-15张)", level=2)
supp_figs = [
    "S1: 两个bulk RNA-seq数据集的QC (PCA, sample correlation, dispersion plots)",
    "S2: GSEA的permutation test结果",
    "S3: WGCNA诊断图 (scale-free fit, module dendrogram, module-trait correlation)",
    "S4: snRNA-seq QC (violin plots of nUMI/nGene, mitochondrial%, doublet scores)",
    "S5: 所有9种细胞类型的marker gene dotplot",
    "S6: CellChat全部交互路径 (非仅髓系->心肌)",
    "S7: NicheNet补充分析 (如果跑了)",
    "S8: GSE57338的完整DEG结果 (火山图+热图)",
    "S9: 三平台蛋白质组的独立mRNA-蛋白散点图 (PXD035154, PXD040135, PXD046631)",
    "S10: OGT-PIN中782个基因在WGCNA模块中的分布",
    "S11: 药物重定位的PAINS过滤过程 (如果保留此部分)",
    "S12: Boolean network robustness测试的多阈值敏感性",
]
for sf in supp_figs: doc.add_paragraph(sf, style="List Bullet")

doc.add_paragraph("")

# ============================================================
doc.add_heading("阶段三：手稿撰写 (Day 8-21)", level=1)

writing_steps = [
    {
        "step": "3.1 STAR Methods (Day 8-9)",
        "target": "另一实验室能根据Methods完全复现你的分析",
        "action": [
            "Key Resources Table (抗体/试剂/软件/数据 accession)",
            "Experimental Model: 小鼠I/R模型参数，样本数，时间点",
            "Method Details: RNA-seq预处理(STAR, featureCounts), DESeq2参数, snRNA-seq(Cell Ranger, Seurat), 蛋白组(MaxQuant), CellChat参数, NicheNet参数, WGCNA参数",
            "Quantification: 统计检验方法，P值校正方法，GSEA参数，Fisher精确检验",
            "Data and Code Availability: GEO/PRIDE accession, GitHub URL",
        ],
        "verify": "让不了解该项目的人阅读Methods后列出所需软件和数据，与实际一致",
    },
    {
        "step": "3.2 Results 第一轮 — 叙事草稿 (Day 10-13)",
        "target": "7个子标题，每段500-800字，全部按'问题->数据->结论'结构",
        "action": [
            "只写文字，不嵌入Figure引用细节（先保证叙事流畅）",
            "每段第一句必须是一个科学问题/假设",
            "每段最后一句是结论+过渡到下一段",
            "写完一段后立即大声朗读——不顺口就改",
        ],
        "verify": "连续朗读全部Results，不卡顿、不跳跃、逻辑自洽。卡住的地方重写。",
    },
    {
        "step": "3.3 Results 第二轮 — 数据锚定 (Day 14-15)",
        "target": "每个claim都有Figure/Table支撑，无空口无凭的陈述",
        "action": [
            "逐句标注Figure引用",
            "检查：每个主Figure在Results中至少被引用2次",
            "补充Statistics: 每个差异表达/富集都报告 effect size + P value",
            "检查：所有数字在文中的报告与Figure中一致",
        ],
        "verify": "打印Results，用荧光笔标出每个Figure引用；未被引用的Figure需删除或补充文字",
    },
    {
        "step": "3.4 Introduction (Day 16)",
        "target": "3段，400字，漏斗结构",
        "action": [
            "Para 1: 心肌I/R的临床重要性 + O-GlcNAc修饰的研究现状",
            "Para 2: Nrf2/铁死亡在I/R中的已知角色 + 知识缺口(Ogt如何连接Nrf2? 细胞类型特异性? 翻译后调控?)",
            "Para 3: 本研究回答的问题 + 多组学策略概述",
        ],
        "verify": "3段读完能回答: (1)为什么重要? (2)别人做了啥? (3)缺什么? (4)我们怎么补?",
    },
    {
        "step": "3.5 Discussion (Day 17-18)",
        "target": "5段，≤1500字，按7.4五段模型",
        "action": [
            "Para 1: 核心发现总结",
            "Para 2: 机制解读(间接调控)",
            "Para 3: 旁分泌放大模型",
            "Para 4: 疾病阶段特异性",
            "Para 5: 局限性与展望",
        ],
        "verify": "Discussion中没有任何一句话是Results中没出现过的(除了限度和展望); 所有新claim都来自本研究数据",
    },
    {
        "step": "3.6 Abstract + Title (Day 19)",
        "target": "Abstract ≤150 words, Title ≤150 chars",
        "action": [
            "Title格式: [核心发现] in [疾病模型] reveals [机制]",
            "Abstract: 1句背景 + 1句方法概述 + 3-4句核心结果 + 1句结论",
        ],
        "verify": "Abstract中每个数字都能在Results中找到出处",
    },
]

for s in writing_steps:
    doc.add_heading(s["step"], level=2)
    p = doc.add_paragraph(); run = p.add_run("目标: "); run.bold = True; p.add_run(s["target"])
    for a in s["action"]: doc.add_paragraph(a, style="List Bullet")
    p = doc.add_paragraph(); run = p.add_run("验证: "); run.bold = True; p.add_run(s["verify"])
    doc.add_paragraph("")

# ============================================================
doc.add_heading("阶段四：打磨与投稿 (Day 22-30)", level=1)

final_steps = [
    {
        "step": "4.1 内部审阅",
        "action": [
            "打印全文 + 所有Figure",
            "第一遍: 只看Figure和Figure Legends，能否独立理解整个故事?",
            "第二遍: 逐段朗读，标记冗长句子(>30 words)和不顺畅处",
            "第三遍: 检查所有数字一致性(文中 vs Figure vs Supplement)",
        ],
        "verify": "三遍审阅后问题清单清零",
    },
    {
        "step": "4.2 格式合规检查",
        "action": [
            "字数统计 (≤7000, 不含Methods/Refs/Legends)",
            "Figure数量和分辨率 (≤7主图, 300dpi min)",
            "Reference格式 (Cell Press superscript style)",
            "STAR Methods完整性",
            "Data Availability声明的accession全部可访问",
        ],
        "verify": "逐项对照Cell Reports Author Guidelines打勾",
    },
    {
        "step": "4.3 Cover Letter",
        "action": [
            "Para 1: 研究和核心发现的1句描述",
            "Para 2: 为什么适合Cell Reports (多组学资源+跨物种验证+机制洞察)",
            "Para 3: 新颖性声明 (急性I/R特异性信号, 非泛化HF标志物)",
            "Para 4: 数据/代码公开声明, 利益冲突声明",
        ],
        "verify": "给不了解该领域的人读Cover Letter，问'你愿意送审吗'",
    },
    {
        "step": "4.4 提交 + Respond to Reviewers 预案",
        "action": [
            "预设审稿人可能的问题(样本量/批次效应/跨物种验证/蛋白覆盖度)",
            "为每个预设问题准备1-2句回复",
        ],
        "verify": "至少准备10个预设问题及其回复",
    },
]

for s in final_steps:
    doc.add_heading(s["step"], level=2)
    for a in s["action"]: doc.add_paragraph(a, style="List Bullet")
    p = doc.add_paragraph(); run = p.add_run("验证: "); run.bold = True; p.add_run(s["verify"])
    doc.add_paragraph("")

# ============================================================
doc.add_heading("执行检查清单", level=1)
checklist = [
    "[ ] 1.1 分析脚本可独立运行",
    "[ ] 1.2 GitHub仓库就绪",
    "[ ] 1.3 所有GEO/PRIDE accession可公开访问",
    "[ ] 2.1 Fig 1 研究概览通过'2分钟测试'",
    "[ ] 2.2 Fig 2 通路+GTEx开关清晰",
    "[ ] 2.3 Fig 3 FTH1/TFRC 10秒可定位",
    "[ ] 2.4 Fig 4 细胞二分一眼可见",
    "[ ] 2.5 Fig 5 SPP1-CD44模式图可独立讲述",
    "[ ] 2.6 Fig 6 C可作为Graphical Abstract",
    "[ ] 2.7 Fig 7 阶段特异性理解清晰",
    "[ ] 3.1 STAR Methods可复现性验证通过",
    "[ ] 3.2 Results叙事大声朗读无卡顿",
    "[ ] 3.3 每个claim都有Figure支撑",
    "[ ] 3.4 Introduction漏斗结构完整",
    "[ ] 3.5 Discussion五段，无Results外的新claim",
    "[ ] 3.6 Abstract <=150 words, Title <=150 chars",
    "[ ] 4.1 三遍内审问题清零",
    "[ ] 4.2 Cell Reports格式合规",
    "[ ] 4.3 Cover Letter通过'愿意送审吗'测试",
    "[ ] 4.4 预设审稿问题 >=10个",
    "[ ] 全文最终朗读一遍，修改所有不顺处",
]
for c in checklist: doc.add_paragraph(c)

# Save
out = r"D:\Cardiac_Ogt_MultiOmics\manuscript\分步执行计划.docx"
doc.save(out)
print(f"Saved: {out}")
