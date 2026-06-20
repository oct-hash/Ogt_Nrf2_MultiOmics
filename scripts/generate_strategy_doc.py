"""Generate strategy document: paper splitting plan + 4-step validation results"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

BASE = Path(r"D:\Cardiac_Ogt_MultiOmics")
doc = Document()

# ── Styles ────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.size = Pt(11)
style.font.name = "Arial"
style.paragraph_format.space_after = Pt(6)

# ── Title ─────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Cardiac_Ogt_MultiOmics — 论文策略与验证报告")
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph(f"生成日期: 2026-06-20")
doc.add_paragraph("")

# ── 1. Overview ───────────────────────────────────────────────────
doc.add_heading("一、论文拆分策略", level=1)

doc.add_heading("论文一（发现篇）：急性 I/R 中 Ogt-Nrf2-铁死亡轴的跨尺度图谱", level=2)
items_1 = [
    "Bulk RNA-seq 差异表达 + GSEA（6,777 DEGs，Nrf2-ARE 与铁死亡通路协同富集 NES=1.77）",
    "snRNA-seq 细胞类型特异性（48,633 核，9 种类型，Ogt 在心肌/周细胞，Nrf2 靶基因在髓系）",
    "三平台蛋白质组验证（LFQ + TMT，mRNA-蛋白一致性 ~50%）",
    "🆕 OGT 底物富集（382/6556 DEGs 为 OGT 互作蛋白，OR=2.02, P<0.001；RELA 是唯一 Nrf2 上游信号中介）",
    "🆕 GTEx v8 健康基线（n=431，Ogt 39.2 TPM vs Hmox1 3.0 TPM，开关模式）",
    "🆕 GSE57338 疾病阶段特异性（177 HF vs 136 NF，Nrf2 靶基因在慢性心衰中反向表达 → 急性 I/R 特异性信号）",
    "🆕 髓系→心肌旁分泌：SPP1-CD44 双方法确认（CellChat ∩ CellPhoneDB，8 对共识交互）",
    "🆕 蛋白层面铁死亡证据：FTH1 背离（铁蛋白自噬），TFRC 维持（铁输入持续）",
]
for item in items_1:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph("投稿目标: Cell Reports (IF=6.9, Q1)")
doc.add_paragraph("")

doc.add_heading("论文二（机制篇）：布尔网络建模确定 BACH1 为 Nrf2 靶基因抑制介导因子及药物重定位", level=2)
items_2 = [
    "27 节点布尔网络（90% 方向一致性，BACH1 为关键抑制节点）",
    "Boolean in silico 扰动（节点敲除/过表达模拟）",
    "分子对接 + 药物重定位（PAINS 过滤）",
    "❌ 需补：独立数据集上的网络方向验证",
    "❌ 需补：多阈值敏感性分析（不同二值化阈值下的网络稳定性）",
    "❌ 需补：对接 top 候选的 MD 模拟精修",
]
for item in items_2:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph("投稿目标: Cell Systems (IF=8.6) 或 Bioinformatics (IF=4.9)")
doc.add_paragraph("")

# ── 2. Storyline ──────────────────────────────────────────────────
doc.add_heading("二、调整后的故事线", level=1)

doc.add_paragraph(
    "核心论点: 在急性心肌 I/R 损伤中，Ogt 下调通过 RELA/NF-κB 间接激活 Nrf2，"
    "髓系细胞通过 SPP1-CD44 旁分泌放大铁死亡信号，铁蛋白自噬释放游离铁驱动心肌细胞铁死亡。"
    "该机制是急性损伤阶段特异性的，在慢性心衰中不保守。"
)

story = [
    ("① 触发", "Ogt 下调（log₂FC = −0.56，padj = 3×10⁻⁷），OGT 底物网络广泛扰动（382/6556 DEGs）"),
    ("② 中介", "Nrf2/铁死亡核心不是 OGT 直接底物 → RELA（NF-κB p65）是唯一上游信号连接点（log₂FC = +1.13）"),
    ("③ 空间格局", "snRNA-seq：Ogt 在心肌/周细胞，Nrf2 靶基因在髓系 → 旁分泌假说"),
    ("④ 旁分泌验证", "SPP1-CD44（CellChat + CellPhoneDB 双方法确认），CD44 是 Nrf2 靶基因"),
    ("⑤ 翻译层面", "mRNA-蛋白 R=0.253，50.7% 方向一致。FTH1 蛋白↓（铁蛋白自噬）、TFRC 蛋白维持（铁输入）、HMOX1/SLC7A11 蛋白低于 MS 检测限"),
    ("⑥ 疾病特异性", "GTEx 健康心脏：Ogt 高、Nrf2 靶基因沉默。GSE57338 慢性心衰：Nrf2 靶基因反向表达 → 急性 I/R 特异性"),
]

for tag, desc in story:
    p = doc.add_paragraph()
    run = p.add_run(f"{tag}: ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph("")

# ── 3. Four Steps Summary ────────────────────────────────────────
doc.add_heading("三、四项验证结果汇总", level=1)

# Step 3
doc.add_heading("Step ③ OGT 底物靶向富集", level=2)
doc.add_paragraph("数据: OGT-PIN (782 人类高置信度 OGT 互作蛋白) × GSE193997 DEGs (6,556, padj<0.05)")
doc.add_paragraph("结果: 382 个重叠基因 (5.83%), Fisher OR=2.02, P<0.000001, 280 UP / 102 DOWN")
doc.add_paragraph("关键: NFE2L2, KEAP1, BACH1, HMOX1, SLC7A11, GPX4 均不在 OGT-PIN 中 → OGT 对 Nrf2/铁死亡的调控是间接的")
doc.add_paragraph("信号中介: RELA/NF-κB p65 是唯一同时属于 OGT 互作组且显著上调的转录因子 (log₂FC=+1.13, padj=5.3×10⁻¹⁵)")

# Step 4
doc.add_heading("Step ④ 外部独立验证", level=2)
doc.add_paragraph("GTEx v8 (n=431 健康左心室): OGT 39.2 TPM, HMOX1 3.0 TPM, SLC7A11 ~0 TPM → 健康开关模式")
doc.add_paragraph("GSE57338 (177 HF vs 136 NF, 人慢性心衰): 方向一致性 31.5% (低于随机), Nrf2 靶基因在慢性心衰中方向反转")
doc.add_paragraph("→ 该信号是急性 I/R 阶段特异性的，非泛化心衰标志物")

# Step 1
doc.add_heading("Step ① 细胞通讯双方法交叉验证", level=2)
doc.add_paragraph("方法: CellChat ∩ CellPhoneDB (LIANA 框架), 9 种细胞类型, GSE240848 snRNA-seq")
doc.add_paragraph("共识: 8 对交互 (髓系→心肌/周细胞), SPP1-CD44 为核心信号 (巨噬+中性粒→心肌+周细胞)")
doc.add_paragraph("关联: CD44 是 Nrf2 靶基因，介导铁内吞 → 连接旁分泌与铁死亡")

# Step 2
doc.add_heading("Step ② mRNA-蛋白偏离分析（残差法）", level=2)
doc.add_paragraph("数据: PXD046631 (2,308 个匹配基因), 线性回归: R=0.253, 斜率=0.144")
doc.add_paragraph("方向一致: 50.7% (1171/2308), 方向背离: 49.3%")
doc.add_paragraph("铁死亡关键: FTH1 背离 (mRNA↑蛋白↓ → 铁蛋白自噬), TFRC 背离 (mRNA↓蛋白→ → 铁输入维持)")
doc.add_paragraph("Nrf2 靶基因 HMOX1, SLC7A11 蛋白低于 MS 检测限 (mRNA 强诱导但蛋白未检出 → 翻译延迟或降解)")

doc.add_paragraph("")

# ── 4. Data sources ──────────────────────────────────────────────
doc.add_heading("四、数据文件清单", level=1)

data_files = [
    ("外部验证", [
        "validation_data/GTEx_core_genes_LV.json — GTEx v8 5基因×54组织",
        "validation_data/GSE57338_series_matrix.txt.gz — GSE57338 313样本微阵列",
        "validation_data/GPL11532.annot.gz — GPL11532 平台注释",
    ]),
    ("OGT 底物", [
        "validation_data/OGT-PIN_S1.xlsx — OGT-PIN 3577条互作记录",
        "validation_data/OGT_high_genes.json — 782个人类高置信度OGT互作蛋白",
    ]),
    ("分析结果", [
        "output_tables/CCC_consensus_CellChat_CellPhoneDB.csv — 8对共识交互",
        "output_tables/GSE57338_direction_consistency.csv — 方向一致性检验",
        "output_tables/GSE57338_GSEA_ranking.rnk — GSEA 排序文件",
        "output_tables/mRNA_protein_residuals.csv — mRNA-蛋白残差",
        "output_tables/GSE57338_direction_consistency.csv — 方向一致性",
    ]),
    ("源数据 (D:\\miri_paper\\)", [
        "results/tables/GSE193997_DESeq2_Sham_vs_IR6h.csv — DEG 列表",
        "results/tables/SuppTable_CellChat_LR_interactions.csv — CellChat 结果",
        "results/tables/LIANA_liana_res_full.csv — LIANA 结果",
        "results/proteomics/PXD046631_mRNA_protein_merged.csv — 蛋白组+转录组",
    ]),
]

for cat, files in data_files:
    doc.add_heading(cat, level=3)
    for f in files:
        doc.add_paragraph(f, style="List Bullet")

doc.add_paragraph("")

# ── 5. Next Steps ─────────────────────────────────────────────────
doc.add_heading("五、下一步行动", level=1)

next_steps = [
    "论文一：将四项新增结果写入 Methods + Results + Discussion，更新 Figure 排版",
    "论文一：撰写 cover letter，强调 (1) 多组学互补验证 (2) 双方法 CCC (3) 跨物种/跨平台外部验证 (4) 疾病阶段特异性的诚实报告",
    "论文二（独立推进）：补布尔网络独立验证 + 多阈值敏感性分析 + MD 模拟",
    "两篇都完成后：论文二引用已接收/已投稿的论文一",
]
for i, step in enumerate(next_steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_paragraph("")

# ── 6. Proposed Figures ──────────────────────────────────────────
doc.add_heading("六、论文一建议图表布局", level=1)

figs = [
    ("Fig 1", "研究概览 + 多组学整合流程图"),
    ("Fig 2", "Bulk RNA-seq: 火山图 + 热图 (6,777 DEGs)"),
    ("Fig 3", "GSEA: Nrf2-ARE + Ferroptosis 协同富集 + GTEx 健康基线对比"),
    ("Fig 4", "snRNA-seq: UMAP + 细胞类型特异性表达 (Ogt vs Nrf2 targets 的空间二分)"),
    ("Fig 5", "CCC: CellChat + CellPhoneDB 共识，SPP1-CD44 为核心"),
    ("Fig 6", "蛋白质组: mRNA-蛋白散点图 + 残差分析 + 铁死亡节点标注"),
    ("Fig 7", "OGT 底物网络: 富集统计 + RELA 信号中介模型"),
    ("Fig 8", "外部验证: GSE57338 方向一致性散点图 + 疾病阶段特异性模式图"),
    ("Fig 9", "整合模型: Ogt → RELA → Nrf2 → 铁死亡, 急性 I/R 特异性"),
]
for fig_id, desc in figs:
    doc.add_paragraph(f"{fig_id}: {desc}", style="List Bullet")

# ── Save ──────────────────────────────────────────────────────────
out_path = BASE / "output_tables" / "论文策略与验证报告.docx"
doc.save(str(out_path))
print(f"Saved: {out_path}")
