"""Append writing guide to strategy document"""
from docx import Document
from docx.shared import Pt
from pathlib import Path

BASE = Path(r"D:\Cardiac_Ogt_MultiOmics")
doc = Document(str(BASE / "output_tables" / "论文策略与验证报告.docx"))

# ── 7. Writing Guide ──────────────────────────────────────────────
doc.add_heading("七、论文一写作指南", level=1)

doc.add_heading("7.1 Cell Reports 投稿要求", level=2)
specs = [
    ("字数上限", "≤7,000 字 (不含 STAR Methods、参考文献、图注)"),
    ("图表上限", "≤7 个主图/表"),
    ("结构", "Title → Summary (≤150 words) → Introduction → Results → Discussion → STAR Methods → Supplemental"),
    ("STAR Methods", "Resource Availability / Experimental Model / Method Details / Quantification & Statistical Analysis"),
    ("数据可用性", "所有组学数据必须存入公共数据库 (GEO, PRIDE/ProteomeXchange)，文中声明 accession number"),
]
for label, content in specs:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(content)

doc.add_heading("7.2 叙事框架：非'数据列表'式写作", level=2)

doc.add_paragraph(
    "多组学论文最致命的错误是按技术平台（RNA-seq → 蛋白组 → snRNA → ...）'
    '逐个报告结果，变成'数据列表'。核心原则：按生物学发现组织叙事，而非按技术平台。"
)

doc.add_paragraph(
    "我们的核心一句话（North Star）：'急性心肌I/R中，Ogt下调通过RELA/NF-κB间接激活Nrf2，'
    '髓系SPP1-CD44旁分泌放大铁死亡信号，铁蛋白自噬加剧心肌损伤——该机制是急性损伤阶段特异性的。'"
)

doc.add_heading("7.3 推荐的 Results 子标题（不应按技术平台命名）", level=2)

sections = [
    ("Multi-omics profiling identifies Ogt downregulation and Nrf2/ferroptosis co-activation in acute cardiac I/R",
     "Fig 1-2: 研究概览 + DEG + GSEA"),
    ("Cross-platform proteomics reveals predominant post-transcriptional regulation and ferroptosis protein-level signatures",
     "Fig 3: mRNA-蛋白残差 + FTH1/TFRC 背离"),
    ("Single-nucleus transcriptomics uncovers cell-type dichotomy: Ogt in cardiomyocytes, Nrf2 effectors in myeloid cells",
     "Fig 4: snRNA-seq UMAP + 细胞类型特异性"),
    ("Dual-method cell-cell communication inference confirms myeloid-derived SPP1-CD44 as a paracrine ferroptosis amplifier",
     "Fig 5: CellChat ∩ CellPhoneDB 共识"),
    ("OGT substrate network perturbation reveals RELA/NF-kB as the signaling intermediary to Nrf2 activation",
     "Fig 6: OGT-PIN 富集 + RELA 模型"),
    ("External validation confirms acute I/R specificity: Nrf2 targets are silenced in healthy hearts and reversed in chronic heart failure",
     "Fig 7: GTEx 基线 + GSE57338 方向对比"),
]

for title, figs in sections:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}")
    run.bold = True
    p.add_run(f"  [{figs}]")

doc.add_heading("7.4 Discussion 五段模型", level=2)

discussion = [
    ("Paragraph 1 — 总结核心发现",
     "简洁重述 Ogt→RELA→Nrf2→铁死亡的调控轴及其急性I/R特异性。不重复Results数据。"),
    ("Paragraph 2 — 机制解读：为什么是间接调控",
     "OGT-PIN 结果：Nrf2/铁死亡核心不是直接底物 → RELA/NF-κB 是信号中介 → O-GlcNAc修饰对NF-κB的已知调控 → 翻译后调控层的重要性（呼应~50% mRNA-蛋白不一致）。"),
    ("Paragraph 3 — 旁分泌放大模型",
     "SPP1-CD44 双方法确认 → CD44 是 Nrf2 靶基因 → 铁内吞 → 铁死亡敏感性。与已有文献中巨噬细胞-心肌细胞通讯的比较。"),
    ("Paragraph 4 — 疾病阶段特异性",
     "GTEx 健康开关 + GSE57338 慢性心衰中Nrf2靶基因反向 → 该轴是急性损伤反应而非慢性重构机制。这一发现的临床意义：时间窗口治疗、生物标志物阶段依赖性。"),
    ("Paragraph 5 — 局限性与展望",
     "① 小鼠I/R模型，需大动物/临床样本验证 ② 蛋白组覆盖度有限，HMOX1/SLC7A11低于MS检测限 ③ 布尔网络建模为后续机制验证提供了路线图（引用论文二） ④ 翻译效率和蛋白降解无法用现有数据区分。"),
]

for title, content in discussion:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}")
    run.bold = True
    p.add_run(f" — {content}")

doc.add_heading("7.5 参考范文", level=2)

refs = [
    ("Cell Reports Medicine (2024)",
     "Perry et al. 'Clinical-transcriptional prioritization of the circulating proteome in human heart failure'",
     "DOI: 10.1016/j.xcrm.2024.101704",
     "多组学心脏论文范式：血浆蛋白组 + snRNA-seq + 临床表型三角验证。学习其 Discussion 如何从组学发现过渡到临床意义。"),
    ("Cell Death Discovery (2025)",
     "Nrf2 transcriptional networks and ferroptosis in lung cancer",
     "DOI: 10.1038/s41420-025-02564-z",
     "Nrf2-铁死亡轴的标准论证路径：差异表达 → GSEA → 关键节点验证。"),
    ("J Cell Mol Med (2025)",
     "ILA inhibits ferroptosis via AhR/Nrf2 in Dox-induced cardiotoxicity — snRNA-seq + bulk RNA-seq",
     "DOI: 10.1111/jcmm.70358",
     "snRNA-seq 在心脏铁死亡中的直接应用。注意其 Figure 布局：先 bulk DEG → 再 snRNA UMAP → 最后通路验证。"),
]

for title, citation, doi, note in refs:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(f"{citation} ({doi}). {note}")

doc.add_heading("7.6 常见写作陷阱", level=2)
pitfalls = [
    "❌ 'Figure 2A shows the volcano plot of DEGs...' → ✅ 'Ogt was significantly downregulated (log₂FC = -0.56, padj = 3×10⁻⁷), accompanied by Nrf2 target gene induction (Fig 2A).'",
    "❌ 'Next, we performed proteomics analysis...' → ✅ 'To determine whether transcriptional changes propagated to the protein level...'",
    "❌ 'Interestingly, HMOX1 was upregulated...' → 去掉'Interestingly'，让读者自己判断什么是有趣的",
    "❌ 在 Results 里写 Discussion 式的推测 → 全部留给 Discussion",
    "❌ 把每个组学作为独立子标题（'Transcriptomic Analysis' / 'Proteomic Analysis'）",
    "❌ 报告所有6,777 DEGs的细节 → 聚焦核心通路，细节放 Supplementary",
    "❌ 忽略阴性结果：GSE57338 中 Nrf2 靶基因反向变化应主动报告，这是论文亮点",
]
for p_text in pitfalls:
    doc.add_paragraph(p_text, style="List Bullet")

doc.add_heading("7.7 GitHub / 代码资源", level=2)
gh = [
    "MOI (MultiOmicsIntegrator): github.com/ASAGlab/MOI — Nextflow 多组学流程，含 DEG、富集、相关性分析模板",
    "OmicVerse: pypi.org/project/omicverse/ — Python 多组学框架，含 Scanpy、CellPhoneDB、WGCNA 整合",
    "rna-seq-proteomics-pipeline: github.com/bmascat/rna-seq-proteomics-pipeline — RNA-seq + LFQ 蛋白组整合模板",
    "建议: 将你的分析脚本整理成一个 GitHub repo，在 STAR Methods 中引用，提升可重复性评价",
]
for g in gh:
    doc.add_paragraph(g, style="List Bullet")

doc.add_heading("7.8 Figure 布局建议 (≤7 主图)", level=2)
figs_layout = [
    ("Fig 1. 研究概览与多组学整合全景",
     "A: 研究设计示意图 (Bulk RNA-seq x2 + snRNA-seq + scRNA-seq + 蛋白组x3 → 分析流程) "
     "B: 火山图 (6,777 DEGs, Ogt/Nrf2 靶基因标注)",
     "要点: 第一张图要让审稿人在2分钟内理解全部研究规模"),
    ("Fig 2. Nrf2-ARE与铁死亡通路协同激活",
     "A: GSEA 双通路富集图 (NES=1.77, FDR=0.001)  B: 热图 (核心基因表达模式)  "
     "C: GTEx 健康基线对比 (Ogt 高 / Nrf2靶基因低 → 开关模式)",
     "要点: 把外部验证 (GTEx) 嵌入此图，增加可信度"),
    ("Fig 3. mRNA-蛋白偏离揭示铁死亡翻译后调控",
     "A: 散点图 (2,308基因, R=0.253)  B: 残差排名 + GSEA 结果  "
     "C: FTH1/TFRC 背离放大图 + 铁蛋白自噬模式图",
     "要点: 从50%一致性落地为具体基因的机制解释"),
    ("Fig 4. snRNA-seq揭示细胞类型空间二分",
     "A: UMAP (48,633核, 9种类型)  B: Ogt vs Nrf2靶基因的Feature Plot  "
     "C: 细胞类型比例 + 关键基因小提琴图",
     "要点: 视觉上要一眼看出'Ogt在心肌、Nrf2在髓系'的空间分离"),
    ("Fig 5. SPP1-CD44: 双方法确认的髓系→心肌旁分泌信号",
     "A: CellChat 互作网络图  B: CellChat ∩ CellPhoneDB 共识Venn/UpSet  "
     "C: SPP1-CD44 信号通路模式图 + 下游 Nrf2/铁死亡连接",
     "要点: 强调双方法确认增加了旁分泌假说的可信度"),
    ("Fig 6. OGT底物网络与RELA信号中介",
     "A: OGT-PIN × DEG 富集统计 (OR=2.02, P<0.001)  B: 27节点网络中的OGT互作关系  "
     "C: RELA 作为唯一信号中介的模式图 (Ogt → RELA/NF-κB → Nrf2 → 铁死亡)",
     "要点: 从'底物富集'到'间接调控模型'的演绎"),
    ("Fig 7. 外部验证确认急性I/R阶段特异性",
     "A: GSE57338 方向一致性散点图 (31.5%, HMOX1/SLC7A11 反向)  "
     "B: 急性I/R vs 慢性心衰的模式对比图  C: 整合模型总结",
     "要点: 诚实报告'不验证'的反而是更高级的科学叙事"),
]

for fig_id, panels, note in figs_layout:
    p = doc.add_paragraph()
    run = p.add_run(f"{fig_id}")
    run.bold = True
    p.add_run(f"\n  Panels: {panels}\n  ⚡ {note}")

doc.add_heading("7.9 时间线建议", level=2)
timeline = [
    "Week 1: 完成 Methods 部分 (STAR Methods 格式)，整理 Supplement",
    "Week 2-3: 撰写 Results (按上述7个子标题)，每个子标题 500-800 字",
    "Week 4: 撰写 Introduction + Discussion",
    "Week 5: 制作 Figure (Python/R → 导出 PDF/EPS)，写 Figure Legends",
    "Week 6: 内部审阅 + 修改 + 格式检查 + Cover Letter",
]
for t in timeline:
    doc.add_paragraph(t, style="List Bullet")

# ── Save ──────────────────────────────────────────────────────────
out_path = BASE / "output_tables" / "论文策略与验证报告.docx"
doc.save(str(out_path))
print(f"Updated: {out_path}")
