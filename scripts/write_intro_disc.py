"""Generate Introduction + Discussion draft"""
from docx import Document; from docx.shared import Pt; from pathlib import Path

doc = Document()
s = doc.styles["Normal"]; s.font.size = Pt(11); s.font.name = "Arial"
s.paragraph_format.space_after = Pt(6)

def h(text, level=1): doc.add_heading(text, level=level)
def p(text): doc.add_paragraph(text)

h("Introduction", level=1)
p("Myocardial ischemia-reperfusion (I/R) injury is a major determinant of infarct size "
  "and clinical outcomes following acute myocardial infarction. Despite advances in "
  "reperfusion therapy, the molecular mechanisms governing cardiomyocyte survival and "
  "death during the early reperfusion phase remain incompletely understood, and targeted "
  "therapies to limit I/R injury have largely failed in clinical translation.")

p("O-GlcNAcylation, a dynamic and reversible post-translational modification catalyzed "
  "by O-GlcNAc transferase (OGT) and removed by O-GlcNAcase (OGA), functions as a "
  "nutrient and stress sensor linking cellular metabolism to signal transduction. In the "
  "heart, O-GlcNAcylation is acutely protective during ischemia, but its role during "
  "reperfusion is complex and context-dependent. Nuclear factor erythroid 2-related factor "
  "2 (Nrf2) is the master transcriptional regulator of the antioxidant response, controlling "
  "genes involved in glutathione synthesis, redox homeostasis, and iron metabolism. "
  "Ferroptosis, an iron-dependent form of regulated cell death driven by lipid peroxidation, "
  "has recently been implicated in cardiac I/R injury. However, the molecular connections "
  "linking O-GlcNAc signaling, Nrf2 activation, and ferroptosis in the injured heart remain "
  "poorly defined. In particular, whether OGT directly or indirectly regulates the "
  "Nrf2/ferroptosis axis, which cell types orchestrate this regulation, and how "
  "transcriptional changes translate to the proteome are unresolved questions.")

p("Here, we present an integrated multi-omics analysis combining bulk RNA-seq, single-nucleus "
  "RNA-seq (snRNA-seq), single-cell RNA-seq (scRNA-seq), and three independent proteomics "
  "datasets to systematically characterize the OGT-Nrf2-ferroptosis axis in acute myocardial "
  "I/R injury. We identify Ogt downregulation as a trigger for Nrf2 pathway activation, "
  "reveal a cell-type dichotomy in which Ogt is enriched in cardiomyocytes and pericytes "
  "while Nrf2 effectors are selectively expressed in myeloid cells, and confirm "
  "myeloid-to-cardiomyocyte SPP1-CD44 paracrine signaling as a ferroptosis-amplifying "
  "circuit. Using cross-platform proteomics, we demonstrate that approximately 50% of "
  "transcriptional changes propagate to the protein level, with ferroptosis-related proteins "
  "showing distinct discordance patterns indicative of active ferritinophagy. A curated OGT "
  "protein interaction network combined with our transcriptomic data identifies RELA/NF-kB "
  "as the signaling intermediary between OGT and Nrf2. Finally, external validation in "
  "healthy (GTEx, n = 431) and chronic heart failure (GSE57338, n = 313) human cohorts "
  "establishes the acute-injury specificity of this regulatory axis.")

h("Discussion", level=1)
p("In this study, we integrated bulk and single-nucleus transcriptomics with three "
  "independent proteomics platforms to systematically dissect the OGT-Nrf2-ferroptosis "
  "regulatory axis in acute myocardial I/R injury. Our findings establish four key "
  "principles: (1) Ogt downregulation indirectly activates Nrf2 through RELA/NF-kB, "
  "rather than through direct substrate interaction; (2) a cell-type dichotomy partitions "
  "Ogt expression to cardiomyocytes and pericytes while confining Nrf2 effector expression "
  "to myeloid cells; (3) myeloid-derived SPP1 signaling through CD44 on cardiomyocytes "
  "constitutes a paracrine ferroptosis-amplifying circuit; and (4) this regulatory axis is "
  "specific to acute I/R injury, not a generic feature of cardiac dysfunction (Fig 7C).")

p("The absence of core Nrf2 and ferroptosis proteins from the OGT interactome (OGT-PIN) "
  "is a central finding of this study. While OGT substrates were globally enriched among "
  "DEGs (OR = 2.01, P < 0.001), NFE2L2, KEAP1, BACH1, HMOX1, SLC7A11, and GPX4 were "
  "completely absent from the curated OGT protein interaction network. This negative result "
  "redirects attention to upstream signaling intermediates, particularly RELA/NF-kB p65, "
  "the sole transcription factor at the intersection of the OGT interactome and the I/R "
  "transcriptomic response. O-GlcNAcylation of RELA at Thr352 is known to enhance NF-kB "
  "transcriptional activity, and NF-kB binding sites are present in the NFE2L2 promoter. "
  "Our data thus support an Ogt-RELA-Nrf2 signaling relay, in which Ogt downregulation "
  "reduces RELA O-GlcNAcylation, modulating NF-kB-dependent Nrf2 transcription.")

p("The cell-type dichotomy observed in our snRNA-seq analysis was an unexpected finding "
  "that prompted cell-cell communication analysis. The convergence of CellChat and "
  "CellPhoneDB on SPP1-CD44 as the dominant myeloid-to-cardiomyocyte signal is particularly "
  "noteworthy given the identity of CD44 as a direct Nrf2 transcriptional target and its "
  "role in iron endocytosis. SPP1 (osteopontin) is a macrophage-derived matricellular "
  "protein implicated in cardiac fibrosis and adverse remodeling. Our data suggest a model "
  "in which Nrf2 activation in myeloid cells drives CD44 surface expression, sensitizing "
  "these cells to autocrine/paracrine SPP1 signaling and thereby coupling inflammatory "
  "activation to iron handling and ferroptosis sensitivity.")

p("The disease-stage specificity of the Nrf2/ferroptosis signature has important "
  "implications. Our cross-platform validation in GSE57338 revealed that key Nrf2 targets "
  "Hmox1 and Slc7a11 move in opposite directions in chronic HF compared to acute I/R. "
  "This finding mirrors the known biphasic role of ferroptosis in cardiac pathology: acute "
  "ferroptosis contributes to infarct expansion, while chronic iron dysregulation promotes "
  "adverse remodeling. It also cautions against interpreting acute-injury transcriptomic "
  "signatures as generic heart failure biomarkers. The stage specificity suggests that "
  "therapeutic windows for O-GlcNAc- or ferroptosis-targeted interventions may be narrow "
  "and injury-phase-dependent.")

p("Several limitations should be noted. First, our transcriptomic data were generated in "
  "mouse I/R models; validation in large-animal models and human acute I/R samples is "
  "needed. Second, the proteomics datasets provided limited coverage; Hmox1 and Slc7a11 "
  "were not detected at the protein level, highlighting the sensitivity gap between "
  "transcriptomics and current proteomics technologies. Third, our analysis cannot "
  "distinguish translational efficiency from protein degradation without ribosome profiling "
  "data. Fourth, the Boolean network modeling and drug repurposing analysis performed in "
  "our companion study (in preparation) provide computational predictions of regulatory "
  "logic that await experimental perturbation. Despite these limitations, the multi-scale, "
  "multi-platform integration presented here establishes the OGT-Nrf2-ferroptosis axis as "
  "an acute-injury-specific regulatory program and provides a foundation for therapeutic "
  "targeting of this pathway in myocardial I/R injury.")

out = Path(r"D:\Cardiac_Ogt_MultiOmics\manuscript\Intro_Discussion_draft.docx")
doc.save(str(out))
print(f"Saved: {out}")
